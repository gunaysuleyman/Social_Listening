import os
import datetime
import pandas as pd
from openai import OpenAI
from docx import Document
from docx.shared import Inches
import time
from dotenv import load_dotenv

load_dotenv()  

def load_test_data():
    """Load data directly from existing Excel files"""
    try:
        
        google_file = "google_trends_data.xlsx"  
        youtube_file = "YouTube_Trend_Videolar.xlsx"  
        
        if not os.path.exists(google_file):
            print(f"Error: {google_file} not found!")
            return None, None
            
        if not os.path.exists(youtube_file):
            print(f"Error: {youtube_file} not found!")
            return None, None
            
        print("Reading Excel files...")
        google_df = pd.read_excel(google_file)
        youtube_df = pd.read_excel(youtube_file)
        
        print("Data loaded successfully!")
        return google_df, youtube_df
        
    except Exception as e:
        print(f"Error loading data: {e}")
        return None, None

def generate_test_report():
    baslangic_zamani = time.time()
    
    
    google_df, youtube_df = load_test_data()
    if google_df is None or youtube_df is None:
        return
 
    print("\nGoogle Trends Data Preview:")
    print(google_df.head())
    print("\nYouTube Trends Data Preview:")
    print(youtube_df.head())
    
  
    data = "GOOGLE TRENDS DATA:\n" + google_df.to_string(index=False)
    data += "\n\nYOUTUBE TRENDS DATA:\n" + youtube_df.to_string(index=False)

    # Verileri trends_data.txt dosyasına kaydet
    #with open("trends_data.txt", "w", encoding="utf-8") as file:
    #    file.write(data)

    #print("Veri trends_data.txt dosyasına kaydedildi.")
    
    
    api_key = os.getenv('OPENAI_API_KEY')
    if not api_key:
        print("Hata: OPENAI_API_KEY bulunamadı! Lütfen .env dosyasını kontrol edin.")
        return
    
    print("\nSending data to GPT-4...")
    client = OpenAI(api_key=api_key)
    
    test_prompt = """You are an expert digital marketing analyst and data scientist. I want you to prepare a comprehensive social media listening report with an influencer marketing focus, based on Google Trends and YouTube Trends data. This report should be in Turkish and cover the following points:

1 DATA SETS 
a) Google Trends Data
    • Search Term / Trend Name
    • Search Volume
    • Start Date
    • Trend Breakdown (context and detailed explanation) 
b) YouTube Trends Data
    • Title
    • Description
    • Upload Date
    • View Count
    • Like Count
    • Comment Count
    • Category (e.g. Music, Film, Gaming, etc.)
2 DATA CLASSIFICATION 
    a) For Google Trends data, classify each search term under the most suitable category:
        • Entertainment (TV, Film, Music, Games)
        • Sports
        • Technology
        • Current Affairs/News
        • Social Media
        • E-commerce/Shopping
        • Health
        • Education
        • Lifestyle
        • Other 
    b) For YouTube Trends data, classify each entry based on the provided category (e.g., Music, Film, Gaming) and content theme.

3 TREND ANALYSIS For each trend (from both Google and YouTube):
    • Describe the core subject or essence of the trend (using Trend Breakdown/Description).
    • Define the target audience (demographics, interests, etc.).
    • Indicate the popularity level (based on Search Volume, View Count, Like Count, etc.).
    • Evaluate social media engagement potential (comments, likes, shares, etc.).
4 DATA ANALYSIS 
    a) Category Distribution
        • Show the number and percentage of Google/YouTube trends in each category. 
    b) Volume/Views Analysis
        • Highlight trends with the highest and lowest figures. 
    c) Time Analysis
        • Indicate distribution according to Start/Upload Dates. 
    d) Target Audience Analysis
        • Identify which demographic segments each trend appeals to.
    5 SOCIAL MEDIA POTENTIAL 
        • Recommend the most suitable social media platforms (Instagram, TikTok, YouTube Shorts, Twitter/X, LinkedIn, etc.) for each trend.
        • Suggest possible content types (short videos, live streams, sponsored posts, post series, etc.) and estimate the expected engagement.
        • Indicate which trends show strong viral potential.
6 INFLUENCER MARKETING FOCUS 
    a) In-depth look at the most popular or high-engagement trends, from an influencer marketing perspective.
    b) Explain key factors in influencer selection (category, content format, brand image match, audience interest, etc.).
    c) Specify which types of influencers or profiles (Music, Lifestyle, Gaming, etc.) would fit each trend category. 
        • Influencer Categories
            • Mega Influencer (1M+ followers)
            • Macro Influencer (100K - 1M followers)
            • Micro Influencer (10K - 100K followers)
            • Nano Influencer (1K - 10K followers) 
    • Recommended Influencer Profile
    • Content category (Lifestyle, Gaming, Beauty, Tech, etc.)
    • Expected engagement rate
    • Platform expertise (Instagram, YouTube, TikTok, etc.)
    • Target audience fit
7 CAMPAIGN STRATEGY RECOMMENDATIONS For each trend:
    • Suggest optimal content format (Reels, Shorts, long-form videos, live streaming, etc.).
    • Recommend platform(s) with justification.
    • Propose campaign duration.
    • Set KPI targets (engagement rate, views, CTR, etc.).
    • Suggest potential collaboration models (co-hosted live streams, sponsored challenges, hashtag campaigns, etc.).
    • Include a risk analysis and mitigation strategies.
8 TREND-INFLUENCER OPPORTUNITIES • Highlight influencer categories matching rising trends.
    • Provide trend-specific content suggestions (e.g., gaming tournaments, acoustic music sessions, etc.).
    • Emphasize seasonal and real-time marketing opportunities.
    • Recommend cross-platform campaign integrations.
9 BRAND OPPORTUNITIES • Identify which brands or brand types could benefit from these trend-influencer collaborations.
    • Offer content ideas and campaign concepts (live events, sponsored challenges, hashtag campaigns, collaboration videos, etc.).
    • Suggest platform strategies (Instagram Reels, YouTube Shorts, TikTok, podcasts, etc.).
10 REPORT FORMAT
    • Executive Summary (key findings and recommendations)
    • Methodology (data sources, analysis methods)
    • Trend Categories and Distribution (both for Google and YouTube)
    • Detailed Trend Analysis
        - Trend descriptions
        - Search volume/view/comment/like analyses
        - Target audience breakdown
    • Social Media Potential
    • Influencer Marketing Opportunities (top influencer categories, matching trends with suitable influencers)
    • Brand Opportunities (content ideas, campaign concepts, platform strategies)
    • Recommendations and Conclusion

11 ADDITIONAL REQUIREMENTS
    • Under a dedicated section “Which Category of Influencers to Work With?”, explicitly match influencer categories (Mega, Macro, Micro, Nano) to each relevant trend category.
    • Provide 5 trending celebrity names and 5 trending music tracks in the report, and explain how they might be integrated into influencer marketing campaigns.
    • Use numerical data, percentages, and comparisons to clarify findings.
    • Include tables where applicable to visualize category distributions, popularity rankings, and other metrics.
    • The report should enable readers to envision how to utilize an influencer marketing system by showing them which influencers to work with and how to collaborate.
    • If any individual mentioned in the analysis has been associated with negative news, recommend against collaborating with them.
    • Any film or series that appears in the report should also be listed at the very end of the report in JSON format. For example:
{
"film_series": [
{"name": "xxx", "type": "tv"},
{"name": "yyy", "type": "movie"}
]
}
(Note: “xxx” and “yyy” are placeholders; please insert actual film/series names from the report content.)
    • Do not include or analyze any names of football (soccer) or basketball teams, even if they appear in the dataset. These team names must not be mentioned or evaluated in any way.

Please present the entire report, including the analysis, tables, and examples, in Turkish. Address every heading and subheading in detail, particularly the influencer categories (Mega, Macro, Micro, Nano) and how they align with the trend categories, as well as the 5 trending celebrity names and 5 trending music tracks. This will help campaign managers develop more concrete and creative ideas when employing an influencer marketing approach.

"""    


    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": test_prompt},
                {"role": "user", "content": f"Data: {data}"}
            ],
            max_tokens=10000,  
            temperature=0.7
        )
        report_content = response.choices[0].message.content
        token_used = response.usage.total_tokens
        
    except Exception as e:
        print(f"Error while communicating with GPT API: {e}")
        return

  
    print("Creating test report document...")
    today = datetime.date.today().strftime("%Y-%m-%d")
    word_file_name = f"test_report_{today}.docx"

    doc = Document()
    doc.add_heading("Test Social Listening Report", level=1)
    doc.add_paragraph(report_content)
    doc.save(word_file_name)

    bitis_zamani = time.time()
    gecen_sure = round(bitis_zamani - baslangic_zamani, 2)
    print(f"\nTest report generated: {word_file_name}")
    print(f"Program completed in {gecen_sure} seconds.")
    print(f"Harcanan Token: {token_used} ")

if __name__ == "__main__":
    generate_test_report()