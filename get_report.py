import os
import datetime
import pandas as pd
from get_googlecsv import download_trends_csv
from get_youtube_trends import get_trending_videos
from openai import OpenAI
from docx import Document
from docx.shared import Inches
import time
from dotenv import load_dotenv

load_dotenv()  

def load_data():
    """Load and prepare both Google Trends and YouTube Trends data"""
    print("Downloading Google Trends data...")
    download_trends_csv(auto_close=True, headless=True)
    
    current_directory = os.getcwd()
    excel_files = [f for f in os.listdir(current_directory) if f.endswith('.xlsx') and 'YouTube' not in f]
    if not excel_files:
        print("No Google Trends Excel file found. Exiting...")
        return None, None
    
    latest_excel = max([os.path.join(current_directory, f) for f in excel_files], key=os.path.getctime)
    print(f"Google Trends file found: {latest_excel}")
    google_df = pd.read_excel(latest_excel)
    
    youtube_file = "YouTube_Trend_Videolar.xlsx"
    if not os.path.exists(youtube_file):
        print("Running YouTube Trends data collection...")
        from get_youtube_trends import categories, API_KEY
        all_videos = []
        for category_name, category_id in categories.items():
            videos = get_trending_videos(API_KEY, category_id=category_id)
            for video in videos:
                video["Kategori"] = category_name
            all_videos.extend(videos)
        youtube_df = pd.DataFrame(all_videos)
        youtube_df.to_excel(youtube_file, index=False)
    else:
        youtube_df = pd.read_excel(youtube_file)
    
    return google_df, youtube_df

def format_data_for_gpt(google_df, youtube_df):
    """Format both datasets for GPT prompt"""
    google_data = "GOOGLE TRENDS DATA:\n" + google_df.to_string(index=False)
    youtube_data = "\n\nYOUTUBE TRENDS DATA:\n" + youtube_df.to_string(index=False)
    return google_data + youtube_data

def create_table_from_markdown(doc, markdown_text):
    """Convert markdown-style tables to Word tables"""
    lines = markdown_text.split('\n')
    table_lines = []
    current_table = []
    
    for line in lines:
        if '|' in line:
           
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            current_table.append(cells)
        elif current_table:
            if len(current_table) > 1:  
               
                table = doc.add_table(rows=len(current_table), cols=len(current_table[0]))
                table.style = 'Table Grid'
                
                
                for i, row in enumerate(current_table):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
                
                
                doc.add_paragraph()
            current_table = []
        
    
    if current_table and len(current_table) > 1:
        table = doc.add_table(rows=len(current_table), cols=len(current_table[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(current_table):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell
        doc.add_paragraph()

def generate_social_listening_report():
    baslangic_zamani = time.time()
    
    
    google_df, youtube_df = load_data()
    if google_df is None or youtube_df is None:
        return
    
    
    data = format_data_for_gpt(google_df, youtube_df)
    
    print("Sending data to GPT-4...")
    client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
    
    prompt = """You are an expert digital marketing analyst and data scientist. I want you to prepare a comprehensive social media listening report with an influencer marketing focus by analyzing the Google Trends and YouTube Trends data I will provide. Your analysis should be detailed, numerical, and in English. Please use the following data sets and the steps below when creating your report:

DATA SETS:

Google Trends Data
Search Term/Trend Name
Search Volume
Start Date
Trend Breakdown (the context and details of the trend)

YouTube Trends Data
Title
Description
Upload Date
View Count
Like Count
Comment Count
Category (e.g., Music, Film, Gaming, etc.)

Please integrate these data sets with an influencer marketing perspective. The analysis should be structured as follows:

DATA CLASSIFICATION

• For the Google Trends data, classify each search term into the most appropriate category among:
Entertainment (TV, Film, Music, Games)
Sports
Technology
Current Affairs/News
Social Media
E-commerce/Shopping
Health
Education
Lifestyle
Other

• For the YouTube Trends data, classify each entry based on its provided category field (e.g. Film, Music, Gaming) and the content theme.

TREND ANALYSIS

• For each trend (from both Google and YouTube data):
Briefly describe the core subject or essence of the trend (using “Trend Breakdown” / “Description” field).
Define the target audience (demographics, interests, etc.).
Specify the popularity level (based on Search Volume, View Count, Likes, etc.).
Evaluate social media engagement potential (comments, likes, shareability, etc.).

DATA ANALYSIS

• Category Distribution: Show the number and percentage of Google/YouTube trends in each category.
• Volume/Views Analysis: Highlight the trends with the highest and lowest figures.
• Time Analysis: Indicate distribution according to Start/Upload Dates.
• Target Audience Analysis: Identify which demographic segments each trend appeals to.

SOCIAL MEDIA POTENTIAL

• Recommend suitable social media platforms for each trend (e.g., Instagram, TikTok, YouTube Shorts, Twitter/X, LinkedIn).
• Propose potential content types (short videos, live streams, sponsored content, posts series, etc.) and give an engagement forecast.
• Indicate trends with strong viral potential.

INFLUENCER MARKETING FOCUS

• Among the most popular or high-engagement trends, analyze them from an influencer marketing perspective.
• Explain key factors in choosing an influencer (category, content format, audience match, brand image, etc.).
• Suggest which influencers or influencer types would fit each category/“label” (Music, Lifestyle, Gaming, etc.).

• Provide guidance for campaign managers, including:

Which influencer strategies to apply for specific trend categories
Which metrics (engagement rate, subscriber count, watch time, etc.) are critical
Which influencer segment best aligns with the brand’s target audience

BRAND OPPORTUNITIES

• Propose which brands or brand types could benefit from these trend/influencer collaborations.
• Offer content ideas and campaign concepts (live events, sponsored challenges, hashtag campaigns, collaboration videos, etc.).
• Suggest platform strategies (Instagram Reels, YouTube Shorts, TikTok, podcasts, etc.).

REPORT FORMAT

Executive Summary (major findings and recommendations)
Methodology (data sources, analysis methods)
Trend Categories and Distribution (for both Google and YouTube)
Detailed Trend Analysis
Trend descriptions
Search volume and view/comment/like analyses
Target audience breakdown
Social Media Potential
Influencer Marketing Opportunities (top influencer fields, matching trends with the right influencers)
Brand Opportunities (content ideas, campaign concepts, platform strategies)
Recommendations and Conclusion

IMPORTANT NOTES:
• All analysis must be presented in English.
• Use numerical data, percentages, and comparisons to clarify findings.
• For visual presentation, consider using tables.
• Specifically highlight influencer marketing insights and decisions to support campaign managers."""    


    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": f"Data: {data}"}
            ],
            max_tokens=4000,
            temperature=0.7
        )
        report_content = response.choices[0].message.content
        token_used = response.usage.total_tokens
        token_completion = response.usage.completion_tokens
        
    except Exception as e:
        print(f"Error while communicating with GPT API: {e}")
        return

    print("Creating Word document with tables...")
    today = datetime.date.today().strftime("%Y-%m-%d")
    word_file_name = f"social_listening_{today}.docx"

    doc = Document()
    doc.add_heading("Social Listening Report", level=1)
    doc.add_heading(f"Report Date: {today}", level=2)
    
    
    create_table_from_markdown(doc, report_content)
    
    
    doc.add_paragraph(f"\nToken Usage Statistics:", style='Heading 2')
    token_table = doc.add_table(rows=2, cols=2)
    token_table.style = 'Table Grid'
    token_table.cell(0, 0).text = 'Total Tokens'
    token_table.cell(0, 1).text = str(token_used)
    token_table.cell(1, 0).text = 'Completion Tokens'
    token_table.cell(1, 1).text = str(token_completion)
    
    doc.save(word_file_name)
    print(f"Report saved as: {word_file_name}")
    
    bitis_zamani = time.time()
    gecen_sure = round(bitis_zamani - baslangic_zamani, 2)
    print(f"Program completed in {gecen_sure} seconds.")

if __name__ == "__main__":
    generate_social_listening_report()